import os                          # Line 1
import uuid                        # Line 2
import pdfplumber                  # Line 3
from docx import Document          # Line 4
import docx2txt 
import mammoth
from pathlib import Path           # Line 5
from config import UPLOAD_FOLDER, ALLOWED_EXTENSIONS, MAX_FILE_SIZE_MB  # Line 6


# FUNCTION 1 — Validate the uploaded file

def validate_file(filename: str, file_size_bytes: int) -> tuple[bool, str]:   # Line 11
    
    # Check file extension
    extension = Path(filename).suffix.lower()    # Line 14
    
    if extension not in ALLOWED_EXTENSIONS:      # Line 16
        return False, f"File type '{extension}' is not allowed. Only PDF, DOCX, DOC, TXT are accepted."
    
    # Check file size
    size_in_mb = file_size_bytes / (1024 * 1024)     # Line 20
    
    if size_in_mb > MAX_FILE_SIZE_MB:                # Line 22
        return False, f"File is too large ({size_in_mb:.1f}MB). Maximum allowed size is {MAX_FILE_SIZE_MB}MB."
    
    return True, "File is valid"                     # Line 25



# FUNCTION 2 — Save file to disk

def save_file(filename: str, file_content: bytes) -> tuple[str, str]:   # Line 31

    extension = Path(filename).suffix.lower()        # Line 33
    unique_filename = f"{uuid.uuid4()}{extension}"   # Line 34
    file_path = os.path.join(UPLOAD_FOLDER, unique_filename)  # Line 35

    with open(file_path, "wb") as f:                 # Line 37
        f.write(file_content)                        # Line 38

    return unique_filename, file_path                # Line 40



# FUNCTION 3 — Extract raw text from file

def extract_text(file_path: str) -> str:             # Line 46

    extension = Path(file_path).suffix.lower()       # Line 48

    if extension == ".pdf":                          # Line 50
        return extract_text_from_pdf(file_path)

    elif extension == ".docx":                       # Line 53
        return extract_text_from_docx(file_path)
    
    elif extension == ".doc":
        return extract_text_from_doc(file_path)

    elif extension == ".txt":                        # Line 56
        return extract_text_from_txt(file_path)

    else:
        return ""                                    # Line 59


# HELPER — PDF text extraction

def extract_text_from_pdf(file_path: str) -> str:   # Line 65
    text = ""                                        # Line 66

    with pdfplumber.open(file_path) as pdf:          # Line 68
        for page in pdf.pages:                       # Line 69
            page_text = page.extract_text()          # Line 70
            if page_text:                            # Line 71
                text += page_text + "\n"             # Line 72

    return text.strip()                              # Line 74



# HELPER — DOCX text extraction

def extract_text_from_docx(file_path: str) -> str:  # Line 80
    doc = Document(file_path)                        # Line 81
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]  # Line 82
    return "\n".join(paragraphs)                     # Line 83


# HELPER — DOC text extraction using mammoth

def extract_text_from_doc(file_path: str) -> str:
    """
    Extract text from .doc files using mammoth.
    mammoth converts Word documents to plain text.
    Works for both .doc and .docx formats.
    No system dependencies needed.
    """
    try:
        # Method 1 — mammoth extract_raw_text
        # extract_raw_text gets plain text without any HTML formatting
        with open(file_path, "rb") as f:          # "rb" = read binary
            result = mammoth.extract_raw_text(f)  # returns a Result object
            text = result.value                    # .value contains the text
            if text and text.strip():
                print(f"mammoth extracted {len(text)} chars from {file_path}")
                return text.strip()

    except Exception as e:
        print(f"mammoth failed: {e}")

    try:
        # Method 2 — fallback to docx2txt
        import docx2txt
        text = docx2txt.process(file_path)
        if text and text.strip():
            print(f"docx2txt fallback extracted {len(text)} chars")
            return text.strip()

    except Exception as e:
        print(f"docx2txt also failed: {e}")

    try:
        # Method 3 — fallback to python-docx
        # Some .doc files are actually .docx in disguise
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        if text.strip():
            print(f"python-docx fallback extracted {len(text)} chars")
            return text.strip()

    except Exception as e:
        print(f"python-docx also failed: {e}")

    print(f"All methods failed for {file_path}")
    return ""
    

# HELPER — TXT text extraction

def extract_text_from_txt(file_path: str) -> str:   # Line 89
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:  # Line 90
        return f.read()                              # Line 91